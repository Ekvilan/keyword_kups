from docx import Document

def extract_text_from_cell(cell, collected_text):
    cell_text = cell.text.strip()
    if cell_text and cell_text not in collected_text:
        collected_text.add(cell_text)
        return cell_text
    return ""

def extract_text_from_table(table, collected_text):
    full_text = ""
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            cell_text = extract_text_from_cell(cell, collected_text)
            row_text.append(cell_text)
            # Обработка вложенных таблиц
            for nested_table in cell.tables:
                full_text += extract_text_from_table(nested_table, collected_text)
        full_text += ' '.join(filter(None, row_text)) + '\n'
    return full_text

def extract_text_from_doc(doc_path):
    doc = Document(doc_path)
    collected_text = set()
    full_text = ""
    for table in doc.tables:
        full_text += extract_text_from_table(table, collected_text)
    return full_text.strip()

# Пример использования:
doc_path = 'D:\\create_program\\mail_in_word\\Выписка за 04.06.2024.docx'
extracted_text = extract_text_from_doc(doc_path)
def remove_extra_spaces(text):
    # Разделяем текст на слова и снова соединяем, убирая лишние пробелы
    return ' '.join(text.split())

# Пример использования:
cleaned_text = remove_extra_spaces(extracted_text)
print(cleaned_text)

