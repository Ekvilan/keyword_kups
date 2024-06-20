from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from src.queries.core import select_dataBase
import os

def export_to_word(data, filename="results.docx"):
    """Экспортирует данные в файл Word в виде таблицы."""
    chet = "28."
    info = "Информация о чрезвычайных ситуациях и происшествиях (ДТП, покушения, убийства, катастрофы и др.)."
    doc = Document()

    # Установка размера страницы А4
    section = doc.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(210)

    table = doc.add_table(rows=1, cols=4)
    table.columns[0].width = Inches(1)  # Column 1: Row number
    table.columns[1].width = Inches(8)    # Column 2: Information
    table.columns[2].width = Inches(16)   # Column 3: КУСП
    table.columns[3].width = Inches(3)    # Column 4: Blank
    table.cell(0, 0).text = ""
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = ""
    table.cell(0, 3).text = ""
    new_row = table.add_row()
    new_row.cells[0].text = chet
    new_row.cells[1].text = info
    for chunk in data:
        for row in chunk:
            new_row.cells[2].text += f"{row[0]}.' '{row[1]}\t\n{row[2]}\n"
    new_row.cells[3].text = " "
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles['Normal']
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.runs[0].font.size = Pt(12)
            paragraph.runs[0].font.name = 'Times New Roman'
    doc.save(filename)
    return os.path.abspath(filename)

