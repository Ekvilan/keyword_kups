import re
from docx import Document
import os

def extract_text_from_cell(cell, collected_text):
    cell_text = cell.text.strip()
    if cell_text and cell_text not in collected_text:
        collected_text.add(cell_text)
        return cell_text
    return ""

def extract_text_from_table(table, collected_text):
    full_text = []  # Используем список для хранения текста
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            cell_text = extract_text_from_cell(cell, collected_text)
            row_text.append(cell_text)
            # Обработка вложенных таблиц
            for nested_table in cell.tables:
                full_text.extend(extract_text_from_table(nested_table, collected_text))  # Добавляем текст из вложенных таблиц в список
        # Добавляем текст строки в список
        full_text.append(' '.join(filter(None, row_text)))
    return full_text

def extract_text_from_doc(doc_path):
    doc = Document(doc_path)
    print(doc_path)
    collected_text = set()
    full_text = []  # Используем список для хранения текста
    for table in doc.tables:
        full_text.extend(extract_text_from_table(table, collected_text))  # Добавляем текст из таблицы в список
    return full_text  # Возвращаем список

def split_text_into_kusp(rez):
    text_list = [val for val in rez if val]
    rez_dict = {}
    curr_kysp = None

    for ind, text in enumerate(text_list):
        # Проверяем, является ли следующий элемент началом новой секции
        if ind < len(text_list) - 1 and re.match(r'^(кусп|крсп)', text_list[ind+1].lower()):
           continue
        elif text.lower().startswith('кусп') or text.lower().startswith('крсп'):
            curr_kysp = text
            rez_dict[curr_kysp] = []
        elif curr_kysp:
            rez_dict[curr_kysp].append(text)

    # Возвращаем словарь, а не печатаем его
    return rez_dict

def find_kups(kusp_list, word_find):
    combined_dict = {}
    for key, value_list in kusp_list.items():
        # Объединяем элементы списка в одну строку, разделяя их пробелом
        combined_dict[key] = ' '.join(value_list)
    
    new_dict = {}
    for key, value in combined_dict.items():
        # Проверяем наличие слова 'дтп' в строке
        if 'дтп' in value:
            if any(term in value for term in ['смерть', 'несовершенолетний','труп', 'взрыв','массовое']):
                new_dict[key] = value
        elif 'труп' in value:
            if any(term in value for term in ['ст. 105', 'ст. 111']):
                new_dict[key] = value
        elif any(term in value for term in word_find):
            new_dict[key] = value
    
    return new_dict

word_find = [' взрыв ', ' бомба ', ' террор ', ' дискредитация ', ' скончался ', ' наезд ', ' сбил ', ' сбит ',
              ' горит ', ' пожар ',' возгорание ', ' убил ', ' убийство', ' смерть ', ' смертельные ', ' массовое '
              , ' насильственные действия сексуального характера ', ' угроза убиства ',
            ' несовершеннолетний ', ' малолетний ']

def find_paragraf(folderPath):
    parts = folderPath.split('\\')
    folder_name = '\\'.join(parts[:-1])
    kusp_list = []
    for filename in os.listdir(folder_name):  
        if filename.endswith('.docx'):
            doc_path = os.path.join(folder_name, filename)
            table_text = extract_text_from_doc(doc_path)
            kups_list = split_text_into_kusp(table_text)
            kusp_dict = find_kups(kups_list, word_find)
            kusp_list.append(kusp_dict) 
    print(kups_list)
    return kusp_list



    
